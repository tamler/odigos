"""Agent tool for creating downloadable artifacts (files for the user)."""

from __future__ import annotations

import logging
import mimetypes
import uuid
from datetime import datetime, timezone
from pathlib import Path

from odigos.db import Database
from odigos.storage import FILES_DIR, ARTIFACTS_DIR
from odigos.tools.base import BaseTool, ToolResult

logger = logging.getLogger(__name__)

# Content type mapping for common extensions
_CONTENT_TYPES = {
    ".csv": "text/csv",
    ".md": "text/markdown",
    ".json": "application/json",
    ".html": "text/html",
    ".txt": "text/plain",
    ".xml": "application/xml",
    ".yaml": "application/x-yaml",
    ".yml": "application/x-yaml",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _write_docx(file_path: Path, content: str) -> None:
    """Convert markdown-ish text to a DOCX file."""
    from docx import Document
    doc = Document()
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    doc.save(str(file_path))


class CreateArtifactTool(BaseTool):
    name = "create_artifact"
    category = "create"
    description = (
        "Create a downloadable file for the user. Use this when the user asks you to "
        "generate a spreadsheet, document, report, data export, or any file they can download. "
        "Provide the filename (with extension) and the file content as a string. "
        "Supported formats: CSV, Markdown, JSON, HTML, TXT, XML, YAML, DOCX. "
        "For DOCX: content is plain text, each paragraph separated by newlines. "
        "Lines starting with # become headings."
    )
    parameters_schema = {
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "Filename with extension (e.g. 'report.csv', 'summary.md', 'data.json')",
            },
            "content": {
                "type": "string",
                "description": "The file content as a string",
            },
        },
        "required": ["filename", "content"],
    }

    def __init__(self, db: Database) -> None:
        self.db = db

    async def execute(self, params: dict) -> ToolResult:
        filename = params.get("filename", "").strip()
        content = params.get("content", "")
        conversation_id = params.get("_conversation_id")

        if not filename:
            return ToolResult(success=False, data="", error="Filename is required")

        # Sanitize filename
        filename = Path(filename).name  # Strip any path components
        if not filename or filename.startswith("."):
            return ToolResult(success=False, data="", error="Invalid filename")

        # Determine content type
        ext = Path(filename).suffix.lower()
        content_type = _CONTENT_TYPES.get(ext) or mimetypes.guess_type(filename)[0] or "application/octet-stream"

        # Create artifact -- write to unified data/files/ directory
        artifact_id = str(uuid.uuid4())
        FILES_DIR.mkdir(parents=True, exist_ok=True)
        file_path = FILES_DIR / f"{artifact_id}_{filename}"

        if ext == ".docx":
            _write_docx(file_path, content)
        else:
            file_path.write_text(content, encoding="utf-8")

        file_size = file_path.stat().st_size

        # Register in database
        now = datetime.now(timezone.utc).isoformat()
        await self.db.execute(
            "INSERT INTO artifacts (id, conversation_id, filename, content_type, file_size, file_path, created_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (artifact_id, conversation_id, filename, content_type, file_size, str(file_path), now),
        )

        logger.info("Created artifact %s: %s (%d bytes)", artifact_id[:8], filename, file_size)

        return ToolResult(
            success=True,
            data=f"Created file: {filename} ({file_size} bytes). The user can download it from the artifacts panel.",
            side_effect={
                "artifact": {
                    "id": artifact_id,
                    "filename": filename,
                    "content_type": content_type,
                    "file_size": file_size,
                    "download_url": f"/api/artifacts/{artifact_id}/download",
                },
            },
        )


class DeleteArtifactTool(BaseTool):
    name = "delete_artifact"
    category = "create"
    description = (
        "Delete a file or image by its filename or artifact ID. "
        "Use this when the user asks to remove, delete, or clean "
        "up a generated image or file."
    )
    parameters_schema = {
        "type": "object",
        "properties": {
            "identifier": {
                "type": "string",
                "description": "Filename (e.g. 'report.csv') or artifact UUID to delete",
            },
        },
        "required": ["identifier"],
    }

    def __init__(self, db: Database):
        self.db = db

    async def execute(self, params: dict) -> ToolResult:
        query = (params.get("identifier") or params.get("filename") or "").strip()
        if not query:
            return ToolResult(
                success=False, data="",
                error="No filename provided",
            )

        # Find by ID or filename
        row = await self.db.fetch_one(
            "SELECT id, filename FROM artifacts "
            "WHERE id = ? OR filename = ?",
            (query, query),
        )
        if not row:
            return ToolResult(
                success=False, data="",
                error=f"File not found: {query}",
            )

        artifact_id = row["id"]
        filename = row["filename"]

        # Delete from disk using unified path resolution
        row_full = await self.db.fetch_one(
            "SELECT file_path FROM artifacts WHERE id = ?", (artifact_id,),
        )
        from odigos.storage import resolve_artifact_path
        resolved = resolve_artifact_path(artifact_id, filename, row_full.get("file_path") if row_full else None)
        if resolved and resolved.exists() and not resolved.is_symlink():
            resolved.unlink()
        # Also clean up legacy artifact directory if it exists
        import shutil
        legacy_dir = ARTIFACTS_DIR / artifact_id
        if legacy_dir.exists():
            shutil.rmtree(legacy_dir)

        # Delete from database
        await self.db.execute(
            "DELETE FROM artifacts WHERE id = ?",
            (artifact_id,),
        )

        return ToolResult(
            success=True,
            data=f"Deleted: {filename}",
        )
